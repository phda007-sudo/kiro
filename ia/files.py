"""
Extracao e analise de arquivos de varias extensoes - tudo local (sem APIs).

Responsavel por:
    - extract_text(data, filename): transformar os bytes de um arquivo em texto.
      Formatos de texto/codigo sao lidos direto; PDF/DOCX/XLSX usam bibliotecas
      opcionais (degradacao graciosa: se nao estiverem instaladas, devolve uma
      nota explicando, sem quebrar).
    - analyze(text): estatisticas + palavras-chave + um resumo extrativo simples
      (escolhe as frases mais representativas, sem gerar texto novo).
    - chunk_text(text): quebra o texto em trechos coerentes que serao indexados
      pelo cerebro (cada trecho vira um item pesquisavel no banco).
"""

from __future__ import annotations

import io
import json
import re
from collections import Counter

from . import text as txt

# Extensoes tratadas como texto puro (decodificadas diretamente).
TEXT_EXTS = {
    "txt", "md", "markdown", "rst", "csv", "tsv", "log", "json", "xml",
    "html", "htm", "yaml", "yml", "ini", "cfg", "conf", "toml", "env",
    "py", "js", "ts", "jsx", "tsx", "java", "c", "cpp", "cc", "h", "hpp",
    "cs", "go", "rb", "php", "rs", "css", "scss", "sql", "sh", "bash",
    "bat", "ps1", "r", "kt", "swift", "scala", "pl", "lua", "dart", "vue",
}


def _decode(data: bytes) -> str:
    """Decodifica bytes tentando utf-8 e caindo para latin-1."""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> tuple[str, str]:
    try:
        from pdfminer.high_level import extract_text as _pdf
        return (_pdf(io.BytesIO(data)) or ""), "pdf"
    except ImportError:
        pass
    except Exception as e:  # noqa: BLE001
        return "", f"falha ao ler PDF: {e}"
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = [(p.extract_text() or "") for p in reader.pages]
        return "\n".join(pages), "pdf"
    except ImportError:
        return "", "PDF sem extrator (instale pdfminer.six ou PyPDF2)"
    except Exception as e:  # noqa: BLE001
        return "", f"falha ao ler PDF: {e}"


def _extract_docx(data: bytes) -> tuple[str, str]:
    try:
        import docx
    except ImportError:
        return "", "DOCX sem extrator (instale python-docx)"
    try:
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts), "docx"
    except Exception as e:  # noqa: BLE001
        return "", f"falha ao ler DOCX: {e}"


def _extract_xlsx(data: bytes) -> tuple[str, str]:
    try:
        import openpyxl
    except ImportError:
        return "", "XLSX sem extrator (instale openpyxl)"
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        out: list[str] = []
        for ws in wb.worksheets:
            out.append(f"# planilha: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                out.append(
                    "\t".join("" if c is None else str(c) for c in row)
                )
        return "\n".join(out), "xlsx"
    except Exception as e:  # noqa: BLE001
        return "", f"falha ao ler XLSX: {e}"


def extract_text(data: bytes, filename: str) -> tuple[str, str]:
    """
    Extrai texto dos bytes de um arquivo.

    Retorna (texto, nota). `nota` descreve o tipo/extrator usado ou explica
    por que nao foi possivel extrair (ex.: binario nao suportado).
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(data)
    if ext == "docx":
        return _extract_docx(data)
    if ext == "xlsx":
        return _extract_xlsx(data)

    if ext == "json":
        raw = _decode(data)
        try:
            obj = json.loads(raw)
            return json.dumps(obj, ensure_ascii=False, indent=2), "json"
        except Exception:  # noqa: BLE001
            return raw, "json (texto bruto)"

    if ext in TEXT_EXTS or ext == "":
        return _decode(data), (ext or "texto")

    # Extensao desconhecida: tenta decodificar; se vier muito "lixo", e binario.
    decoded = _decode(data)
    if not decoded or decoded.count("\ufffd") > max(10, len(decoded) * 0.1):
        return "", f"formato '{ext}' nao suportado (binario)"
    return decoded, ext


# --------------------------------------------------------------- analise
def _sentences(text: str) -> list[str]:
    """Divide o texto em frases de forma simples (por pontuacao/linhas)."""
    parts = re.split(r"(?<=[.!?])\s+|\n+", text)
    return [p.strip() for p in parts if p.strip()]


def analyze(text: str, top_n: int = 12, summary_sentences: int = 3) -> dict:
    """
    Analisa um texto e devolve estatisticas, palavras-chave e um resumo
    extrativo (as frases mais representativas, escolhidas pela frequencia
    das palavras-chave que elas contem).
    """
    tokens = txt.tokenize(text)
    counter = Counter(tokens)
    top_keywords = counter.most_common(top_n)

    sentences = _sentences(text)
    resumo = ""
    if sentences:
        freq = counter
        scored: list[tuple[int, float, str]] = []
        for idx, s in enumerate(sentences):
            s_tokens = txt.tokenize(s)
            if not s_tokens:
                continue
            score = sum(freq.get(t, 0) for t in s_tokens) / (len(s_tokens) ** 0.5)
            scored.append((idx, score, s))
        scored.sort(key=lambda x: x[1], reverse=True)
        escolhidas = sorted(scored[:summary_sentences], key=lambda x: x[0])
        resumo = " ".join(s for _, _, s in escolhidas)

    return {
        "caracteres": len(text),
        "linhas": text.count("\n") + 1 if text else 0,
        "palavras": len(text.split()),
        "palavras_unicas": len(counter),
        "palavras_chave": [{"palavra": w, "freq": c} for w, c in top_keywords],
        "resumo": resumo,
    }


def chunk_text(text: str, max_words: int = 35, max_chunks: int = 400) -> list[str]:
    """
    Quebra o texto em trechos coerentes (~max_words palavras), agrupando frases.
    Cada trecho sera indexado como um item pesquisavel no banco.

    Trechos menores melhoram a recuperacao de fatos especificos: uma pergunta
    curta tende a casar melhor com um trecho focado do que com um paragrafo
    longo (a similaridade do cosseno fica menos diluida).
    """
    chunks: list[str] = []
    atual: list[str] = []
    n_palavras = 0

    for s in _sentences(text):
        w = len(s.split())
        if n_palavras + w > max_words and atual:
            chunks.append(" ".join(atual))
            atual, n_palavras = [], 0
            if len(chunks) >= max_chunks:
                return chunks
        atual.append(s)
        n_palavras += w

    if atual and len(chunks) < max_chunks:
        chunks.append(" ".join(atual))

    # Fallback: texto sem pontuacao -> divide por blocos de palavras.
    if not chunks and text.strip():
        words = text.split()
        for i in range(0, len(words), max_words):
            chunks.append(" ".join(words[i : i + max_words]))
            if len(chunks) >= max_chunks:
                break

    return chunks
