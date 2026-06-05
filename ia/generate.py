"""
Geracao de arquivos a partir do conhecimento da IA - tudo local (sem APIs).

Dado um assunto, a IA reune o que ja aprendeu (trechos mais relevantes do
banco) e renderiza um documento no formato/extensao que o usuario pedir:

    - texto/codigo: txt, md, html, json, csv, py, js, java, c, sql, sh, ...
    - binarios: pdf (fpdf2), docx (python-docx), xlsx (openpyxl)

Para formatos binarios as bibliotecas sao opcionais: se faltarem, devolve um
erro claro em vez de quebrar.
"""

from __future__ import annotations

import io
import json
import re
import time

# Tipos MIME por extensao (para o navegador tratar o download corretamente).
MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "html": "text/html; charset=utf-8",
    "htm": "text/html; charset=utf-8",
    "json": "application/json; charset=utf-8",
    "csv": "text/csv; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "txt": "text/plain; charset=utf-8",
}

# Token de comentario de uma linha por extensao de codigo.
COMMENT_TOKEN = {
    "js": "//", "ts": "//", "jsx": "//", "tsx": "//", "java": "//",
    "c": "//", "cpp": "//", "cc": "//", "h": "//", "hpp": "//", "cs": "//",
    "go": "//", "rs": "//", "php": "//", "swift": "//", "kt": "//",
    "scala": "//", "dart": "//",
    "sh": "#", "bash": "#", "rb": "#", "pl": "#", "r": "#", "lua": "--",
    "sql": "--", "yaml": "#", "yml": "#", "toml": "#", "ini": ";",
}


def slugify(texto: str, default: str = "documento") -> str:
    """Transforma o assunto num nome de arquivo seguro."""
    import unicodedata

    t = unicodedata.normalize("NFKD", texto.lower())
    t = "".join(c for c in t if not unicodedata.combining(c))
    t = re.sub(r"[^a-z0-9]+", "_", t).strip("_")
    return (t[:50] or default)


def _latin(s: str) -> str:
    """Adapta o texto para os fontes padrao do PDF (latin-1 cobre o portugues)."""
    return s.encode("latin-1", "replace").decode("latin-1")


# ---------------------------------------------------------- renderizadores
def _render_txt(subject: str, sections: list[str]) -> bytes:
    linhas = [subject, "=" * len(subject), ""]
    linhas += sections if sections else ["(a IA ainda nao tem conhecimento sobre este assunto)"]
    return ("\n\n".join(linhas) + "\n").encode("utf-8")


def _render_md(subject: str, sections: list[str]) -> bytes:
    out = [f"# {subject}", ""]
    if sections:
        for s in sections:
            out.append(f"- {s}")
    else:
        out.append("> A IA ainda nao tem conhecimento sobre este assunto.")
    out += ["", "---", f"_Gerado pela PHDA CEREBROZ em {time.strftime('%d/%m/%Y %H:%M')}_"]
    return ("\n".join(out) + "\n").encode("utf-8")


def _render_html(subject: str, sections: list[str]) -> bytes:
    def esc(s: str) -> str:
        return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    itens = "".join(f"<li>{esc(s)}</li>" for s in sections) or \
        "<li><em>A IA ainda nao tem conhecimento sobre este assunto.</em></li>"
    html = (
        "<!doctype html><html lang=pt-br><head><meta charset=utf-8>"
        f"<title>{esc(subject)}</title></head><body>"
        f"<h1>{esc(subject)}</h1><ul>{itens}</ul>"
        f"<hr><small>Gerado pela PHDA CEREBROZ em {time.strftime('%d/%m/%Y %H:%M')}</small>"
        "</body></html>"
    )
    return html.encode("utf-8")


def _render_json(subject: str, sections: list[str]) -> bytes:
    payload = {
        "assunto": subject,
        "gerado_em": time.strftime("%Y-%m-%d %H:%M:%S"),
        "conteudo": sections,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def _render_csv(subject: str, sections: list[str]) -> bytes:
    import csv

    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["assunto", "trecho"])
    if sections:
        for s in sections:
            w.writerow([subject, s])
    else:
        w.writerow([subject, "(sem conhecimento)"])
    return buf.getvalue().encode("utf-8")


def _render_py(subject: str, sections: list[str]) -> bytes:
    body = "\n\n".join(sections) if sections else "(sem conhecimento sobre o assunto ainda)"
    body = body.replace('"""', '\\"\\"\\"')
    content = (
        '"""\n'
        f"Assunto: {subject}\n"
        "Gerado automaticamente pela PHDA CEREBROZ "
        "(a partir do conhecimento acumulado).\n"
        '"""\n\n'
        f"ASSUNTO = {subject!r}\n\n"
        f'CONTEUDO = """{body}"""\n\n'
        'if __name__ == "__main__":\n'
        '    print("Assunto:", ASSUNTO)\n'
        "    print(CONTEUDO)\n"
    )
    return content.encode("utf-8")


def _render_code(subject: str, sections: list[str], token: str) -> bytes:
    out = [f"{token} === {subject} ===",
           f"{token} Gerado pela PHDA CEREBROZ."]
    corpo = sections if sections else ["(sem conhecimento sobre o assunto ainda)"]
    for s in corpo:
        for linha in s.splitlines() or [s]:
            out.append(f"{token} {linha}")
        out.append(token)
    return ("\n".join(out) + "\n").encode("utf-8")


def _render_pdf(subject: str, sections: list[str]) -> bytes:
    try:
        from fpdf import FPDF
    except ImportError as e:  # noqa: F841
        raise RuntimeError("para gerar PDF instale o pacote 'fpdf2'")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, _latin(subject))
    pdf.ln(2)
    pdf.set_font("Helvetica", size=12)
    corpo = sections if sections else ["A IA ainda nao tem conhecimento sobre este assunto."]
    for s in corpo:
        pdf.multi_cell(0, 7, _latin(s))
        pdf.ln(2)
    pdf.set_font("Helvetica", "I", 9)
    pdf.multi_cell(0, 6, _latin(f"Gerado pela PHDA CEREBROZ em {time.strftime('%d/%m/%Y %H:%M')}"))
    return bytes(pdf.output())


def _render_docx(subject: str, sections: list[str]) -> bytes:
    try:
        import docx
    except ImportError:
        raise RuntimeError("para gerar DOCX instale o pacote 'python-docx'")

    d = docx.Document()
    d.add_heading(subject, level=1)
    if sections:
        for s in sections:
            d.add_paragraph(s)
    else:
        d.add_paragraph("A IA ainda nao tem conhecimento sobre este assunto.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _render_xlsx(subject: str, sections: list[str]) -> bytes:
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("para gerar XLSX instale o pacote 'openpyxl'")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "conteudo"
    ws.append(["assunto", subject])
    ws.append([])
    ws.append(["trecho"])
    for s in (sections or ["(sem conhecimento)"]):
        ws.append([s])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def render(subject: str, sections: list[str], ext: str) -> tuple[str, bytes, str]:
    """
    Renderiza o documento no formato pedido.

    Retorna (nome_do_arquivo, bytes, mimetype).
    """
    ext = (ext or "txt").lower().lstrip(".").strip() or "txt"
    subject = subject.strip()

    if ext == "pdf":
        data = _render_pdf(subject, sections)
    elif ext == "docx":
        data = _render_docx(subject, sections)
    elif ext == "xlsx":
        data = _render_xlsx(subject, sections)
    elif ext == "py":
        data = _render_py(subject, sections)
    elif ext in ("md", "markdown"):
        data = _render_md(subject, sections)
    elif ext in ("html", "htm"):
        data = _render_html(subject, sections)
    elif ext == "json":
        data = _render_json(subject, sections)
    elif ext in ("csv", "tsv"):
        data = _render_csv(subject, sections)
    elif ext in COMMENT_TOKEN:
        data = _render_code(subject, sections, COMMENT_TOKEN[ext])
    else:
        # Qualquer outra extensao: gera como texto simples com aquela extensao.
        data = _render_txt(subject, sections)

    mime = MIME.get(ext, "application/octet-stream")
    filename = f"{slugify(subject)}.{ext}"
    return filename, data, mime
